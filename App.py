from docx.api import Document
from docx.shared import Pt
import sys
import os


global markdown_content
global export_file_name


print("Ensure program folder contains demo.docx and delete README.md to avoid confusion ")

def checkfile(doc_path):
    return os.path.isfile(doc_path)

# def exportFile:
#     retun export_file_name

def checkArguments():
    args=[]
    if(len(sys.argv)>0):
        args=sys.argv
    
    return args

def initialize():
    doc=None
    args=checkArguments()
    if(len(args)==2):
        print(args[1])
        if(os.path.isfile(args[1])):

            doc=Document(str(args[1]))
            print("Input file is set to ",args[1])
        else:
            print("No input file found named ",args[1])
            print("Please check file is present or not and retry again")
            exit()
    
    elif(len(args)==3):
        doc=Document(str(args[1]))
        export_file_name=args[2]
    
    elif(len(args)>3):
        print("Too many arguments")

    else:
        if(checkfile('demo.docx')):
            doc=Document('demo.docx')
            print("converting default file into ReadMe.md")
        else:
            print('end of the line')
            print('retry with a new demo.docx file in root folder of this program or \n paste a docx file and re-run App.py with argument folder name')
            exit()
    
    return doc

def addMainHeading(paragraph):
    return "# "+paragraph.text+"\n\n"

def addHeading(paragraph):
    return "## "+paragraph.text+"\n"

def checkBullets(paragraph):
    flag =False
    if(paragraph.text.strip().startswith('*')):
        flag=True
    else:
        flag=False
    return flag

def formateList(paragraph):
    temp=paragraph.text
    if(temp.strip().startswith("* " or "*")):
        temp=temp.replace("*","")
        temp="* "+temp+"\n"
    elif(paragraph.style.name=="List Bullet"):
        temp=paragraph.text+"\n"
    return temp

def saveMd(content):
    try:
        if(len(sys.argv)==3):
            file_output=open(sys.argv[2],'w')
            file_output.write(markdown_content)
            file_output.close()
            print("Write readme file success")
        else:
            file_output=open(export_file_name,'w')
            file_output.write(markdown_content)
            file_output.close()
            print("Write readme file success")
    except Exception as e:
        print(e)
    
def checkHeading(run):
    flag=False
    if(run.font.size==Pt(14)):
        flag=True
    else:
        flag=False
    return flag

def format_md_content(doc):
    heading =False
    main_heading=False
    for paragraphs in doc.paragraphs:
        # print(paragraphs.text)
        for run in paragraphs.runs:
            if(checkHeading(run)):
                heading=True
            else:
                heading=False
        
        if heading:
            if main_heading:
                markdown_content+=addHeading(paragraphs)
            else:
                markdown_content=addMainHeading(paragraphs)
                main_heading=True
        else:
            # markdown_content+=paragraphs.text+"\n"
            if checkBullets(paragraphs):
                markdown_content+=formateList(paragraphs)
            else:
                markdown_content+=paragraphs.text+"\n"

    # print(markdown_content)
    return markdown_content


def __init__(self):
    print("App is started running")
    self.format_md_content()
    

markdown_content=format_md_content(initialize())
# print(markdown_content)
saveMd(markdown_content)